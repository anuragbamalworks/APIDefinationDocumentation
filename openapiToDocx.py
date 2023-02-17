import yaml
import json
from docx import Document
from docx.enum.text import WD_UNDERLINE
from docx.shared import Pt, Inches
from docx.enum.section import WD_SECTION_START
import jsonschema


def resolve_refs(data,openapi):
    if isinstance(data, dict):
        if '$ref' in data:
            ref = data['$ref']
            if ref.startswith('#/components'):
                parts = ref.split('/')
                obj = openapi
                for part in parts[1:]:
                    obj = obj[part]
                return obj
        else:
            for key in data:
                data[key] = resolve_refs(data[key],openapi)
                    
    elif isinstance(data,list):
        for i in range(len(data)):
            data[i]=resolve_refs(i,openapi)
    return data

def generateApiDocumentaion(yamlFile,docxFile):
    with open(yamlFile, 'r') as file:
        api_spec = yaml.load(file, Loader=yaml.FullLoader)
        resolver = jsonschema.RefResolver.from_schema(api_spec)
    # Create a new Word document
    document = Document()
    #Setting the page size
    section = document.sections[0]
    section.page_height = Inches(30)
    section.page_width = Inches(15)

    # Add title, version, and description to the document
    document.add_heading(api_spec['info']['title'], 0)
    document.add_heading('Version: ' + api_spec['info']['version'], level=1)
    document.add_paragraph(api_spec['info']['description'])

    # Add a section for each API path
    for path, methods in api_spec['paths'].items():
        # Add the path as a section heading
        document.add_heading('Path: ' + path, level=1)

        for method, details in methods.items():
            # Add the HTTP method as a subheading
            document.add_heading('Method: ' + method.upper(), level=2)

            # Add a table to display the query parameters
            if 'parameters' in details and details['parameters']:
                document.add_heading('Query Parameters:', level=3)
                table = document.add_table(rows=1, cols=4)
                table.style = 'Table Grid'
                hdr_cells = table.rows[0].cells
                hdr_cells[0].text = 'Name'
                hdr_cells[1].text = 'Location'
                hdr_cells[2].text = 'Required'
                hdr_cells[3].text = 'Schema'
                for parameter in details['parameters']:
                    row_cells = table.add_row().cells
                    row_cells[0].text = parameter['name']
                    row_cells[1].text = parameter['in']
                    row_cells[2].text = str(parameter['required'])
                    row_cells[3].text = json.dumps(parameter['schema'], indent=4)

            # Add a table to display the response schema
            if 'responses' in details:
                document.add_heading('Response Schema:', level=3)
                table = document.add_table(rows=1, cols=2)
                table.style = 'Table Grid'
                hdr_cells = table.rows[0].cells
                hdr_cells[0].text = 'Status Code'
                hdr_cells[1].text = 'Schema'
                for status_code, response in details['responses'].items():
                    row_cells = table.add_row().cells
                    row_cells[0].text = status_code
                    try:
                        row_cells[1].text = json.dumps(resolve_refs(response['content']['application/json']['schema'],api_spec), indent=8)
                    except Exception as err:
                        print(json.dumps(response['content']))
                        print(path)
                        print(err)

        document.add_page_break()
    # Save the document
    document.save(docxFile)
